"""
Module de parsing de CV automatique.
Phase 1 : Regex + heuristiques (rapide, local, gratuit).
Phase 2 (futur) : Ollama pour les CV complexes.
"""
import re
import io
import base64
import logging
import pdfplumber
from docx import Document
import fitz  # pymupdf - pour extraire les images du PDF
import json
import os
from groq import Groq
from .constants import WILAYAS_MAPPING, DIPLOMES_MAPPING

logger = logging.getLogger(__name__)




# ==========================================
# 1. EXTRACTION DU TEXTE BRUT
# ==========================================
def extract_text_from_pdf(file_path):
    """Extrait tout le texte d'un PDF."""
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        logger.error("Erreur extraction PDF : %s", e)
    return text


def extract_text_from_docx(file_path):
    """Extrait tout le texte d'un fichier Word .docx."""
    text = ""
    try:
        doc = Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
    except Exception as e:
        logger.error("Erreur extraction DOCX : %s", e)
    return text
def extract_text_from_pdf_smart(file_path):
    """
    Extraction intelligente : utilise les coordonnées des mots
    pour mieux comprendre la structure des CV en colonnes.
    Reconstruit les lignes en regroupant les mots qui sont sur la même ligne Y.
    """
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                words = page.extract_words(
                    keep_blank_chars=False,
                    use_text_flow=False,
                    extra_attrs=["fontname", "size"]
                )
                if not words:
                    continue

                # Regroupe les mots par ligne (même position Y approximative)
                lines_dict = {}
                tolerance = 3  # tolérance en pixels pour considérer "même ligne"
                for w in words:
                    y_key = round(w['top'] / tolerance) * tolerance
                    if y_key not in lines_dict:
                        lines_dict[y_key] = []
                    lines_dict[y_key].append(w)

                # Trie chaque ligne par position X (de gauche à droite)
                sorted_y = sorted(lines_dict.keys())
                for y in sorted_y:
                    line_words = sorted(lines_dict[y], key=lambda w: w['x0'])
                    line_text = " ".join(w['text'] for w in line_words)
                    text += line_text + "\n"
    except Exception as e:
        logger.error("Erreur extraction PDF smart : %s", e)
    return text

def extract_text(file_path, file_name):
    """Détecte le type et extrait le texte."""
    name_lower = file_name.lower()
    if name_lower.endswith(".pdf"):
        # On essaie d'abord l'extraction "smart" qui gère les colonnes
        text_smart = extract_text_from_pdf_smart(file_path)
        # Si on a bien du texte, on l'utilise
        if text_smart and len(text_smart.strip()) > 100:
            return text_smart
        # Sinon fallback sur l'extraction classique
        return extract_text_from_pdf(file_path)
    elif name_lower.endswith(".docx") or name_lower.endswith(".doc"):
        return extract_text_from_docx(file_path)
    return ""


# ==========================================
# 2. EXTRACTION DE LA PHOTO DEPUIS LE PDF
# ==========================================
def extract_photo_from_pdf(file_path):
    """
    Extrait la plus grande image du PDF (probablement la photo de profil).
    Retourne une string base64 prête à envoyer au frontend.
    """
    try:
        doc = fitz.open(file_path)
        biggest_image = None
        biggest_size = 0

        for page_num in range(min(2, len(doc))):  # Cherche dans les 2 premières pages
            page = doc[page_num]
            images = page.get_images(full=True)
            for img in images:
                xref = img[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                image_ext = base_image["ext"]

                # On garde la plus grande (généralement la photo de profil)
                if len(image_bytes) > biggest_size and len(image_bytes) > 5000:
                    biggest_size = len(image_bytes)
                    biggest_image = {
                        "data": base64.b64encode(image_bytes).decode('utf-8'),
                        "ext": image_ext,
                    }
        doc.close()
        return biggest_image
    except Exception as e:
        logger.error("Erreur extraction image : %s", e)
        return None


def extract_photo_from_docx(file_path):
    """
    Extrait la plus grande image embarquée dans le fichier Word (probablement la photo de profil).
    Retourne une string base64 prête à envoyer au frontend.
    """
    try:
        doc = Document(file_path)
        biggest_image = None
        biggest_size = 0
        for rel in doc.part.rels.values():
            if rel.is_external or "image" not in rel.reltype:
                continue
            image_part = rel.target_part
            image_bytes = image_part.blob
            if len(image_bytes) > biggest_size and len(image_bytes) > 5000:
                ext = image_part.content_type.split("/")[-1]
                if ext == "jpg":
                    ext = "jpeg"
                biggest_size = len(image_bytes)
                biggest_image = {
                    "data": base64.b64encode(image_bytes).decode('utf-8'),
                    "ext": ext,
                }
        return biggest_image
    except Exception as e:
        logger.error("Erreur extraction image DOCX : %s", e)
        return None


# ==========================================
# 3. CHAMPS SIMPLES (REGEX)
# ==========================================
def extract_email(text):
    pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    match = re.search(pattern, text)
    return match.group(0).lower() if match else None


def extract_phone(text):
    """
    Cherche un numéro algérien peu importe le regroupement de chiffres utilisé
    dans le CV (2-2-2-2-2, 3-3-3, espaces/points/tirets/parenthèses...).
    On capture une plage large de chiffres/séparateurs puis on valide la longueur
    une fois les séparateurs retirés, plutôt que d'imposer un regroupement fixe.
    """
    pattern = r'(?:\(?\+213\)?|00213|0)[\s.\-]?[567][\d\s.\-]{6,14}'
    for match in re.finditer(pattern, text):
        digits = re.sub(r'\D', '', match.group(0))
        if digits.startswith('00213'):
            digits = '0' + digits[5:]
        elif digits.startswith('213') and len(digits) >= 12:
            digits = '0' + digits[3:]
        if len(digits) == 10 and digits[0] == '0':
            return digits
    return None


def extract_name(text):
    """Heuristique : le nom est dans les premières lignes."""
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    for line in lines[:8]:
        if "@" in line or re.search(r'\d{4,}', line):
            continue
        if len(line) > 60 or len(line) < 4:
            continue
        words = line.split()
        if 2 <= len(words) <= 4 and all(len(w) >= 2 for w in words):
            blacklist = ['cv', 'curriculum', 'vitae', 'profil', 'profile']
            if not any(b in line.lower() for b in blacklist):
                return line.title()
    return None


def extract_titre_professionnel(text):
    """
    Cherche le titre pro : c'est souvent la 2e ligne (après le nom),
    en majuscules ou contenant "ingénieur", "développeur", "chef", etc.
    """
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    keywords = [
        'ingénieur', 'ingenieur', 'développeur', 'developpeur', 'chef',
        'directeur', 'responsable', 'manager', 'consultant', 'analyste',
        'cadre', 'technicien', 'assistant', 'comptable', 'commercial',
        'administrateur', 'gestionnaire', 'designer', 'architecte',
        'data scientist', 'expert', 'spécialiste', 'specialiste', 'coordinateur'
    ]
    # Acronymes à garder en majuscules après title()
    acronymes = ['IA', 'IT', 'RH', 'BTP', 'QSE', 'HSE', 'R&D', 'SQL', 'CEO', 'CTO', 'CFO', 'PDG', 'DSI', 'ERP', 'CRM', 'SEO', 'SEM', 'API', 'CV', 'PHP', 'CSS', 'HTML', 'PMO']

    for line in lines[:10]:
        line_low = line.lower()
        if "@" in line or re.search(r'\d{4,}', line):
            continue
        if any(k in line_low for k in keywords) and 5 < len(line) < 120:
            # Capitalise proprement
            result = line.title().replace(" Et ", " et ").replace(" De ", " de ").replace(" Du ", " du ").replace(" La ", " la ").replace(" Le ", " le ").replace(" Des ", " des ")
            # Remet les acronymes en majuscules
            for acro in acronymes:
                # On remplace "Ia" par "IA", "Rh" par "RH", etc. (avec word boundaries)
                pattern = r'\b' + acro.capitalize() + r'\b'
                result = re.sub(pattern, acro, result)
            return result
    return None


def extract_wilaya(text):
    """Détecte la wilaya mentionnée dans le CV."""
    text_low = text.lower()
    for wilaya_text, code in WILAYAS_MAPPING.items():
        # On cherche le mot entier
        pattern = r'\b' + re.escape(wilaya_text) + r'\b'
        if re.search(pattern, text_low):
            return code
    return None


def extract_diplome_max(text):
    """Détecte le plus haut diplôme."""
    text_low = text.lower()
    # On parcourt dans l'ordre (du plus élevé au plus bas)
    for keywords, code in DIPLOMES_MAPPING:
        for kw in keywords:
            if re.search(r'\b' + re.escape(kw) + r'\b', text_low):
                return code
    return None


def extract_specialite(text):
    """Détecte la spécialité principale (code Domaine ANEM) via le référentiel métiers."""
    from .referentiel_utils import resoudre_domaine_depuis_texte
    return resoudre_domaine_depuis_texte(text)


def extract_service_militaire(text):
    """Détecte le statut militaire."""
    text_low = text.lower()
    mapping = {
        'dégagé': 'DEGAGE', 'degage': 'DEGAGE', 'exempted': 'DEGAGE', 'discharged': 'DEGAGE', 'معفى': 'DEGAGE',
        'sursitaire': 'SURSITAIRE', 'sursis': 'SURSITAIRE', 'deferred': 'SURSITAIRE', 'مؤجل': 'SURSITAIRE',
        'inapte': 'INAPTE', 'unfit': 'INAPTE', 'غير لائق': 'INAPTE',
        'incorporé': 'INCORPORE', 'incorpore': 'INCORPORE', 'active duty': 'INCORPORE', 'ملتحق': 'INCORPORE',
        'non concerné': 'NON_CONCERNE', 'non concerne': 'NON_CONCERNE', 'not applicable': 'NON_CONCERNE', 'غير معني': 'NON_CONCERNE',
    }
    for keyword, code in mapping.items():
        if re.search(r'\b' + re.escape(keyword) + r'\b', text_low):
            return code
    return None


def extract_permis(text):
    """Détecte la présence d'un permis de conduire."""
    text_low = text.lower()
    keywords = ['permis de conduire', 'permis b', 'permis a', 'driving license', 'driver license', "driver's license", 'رخصة سياقة', 'رخصة قيادة']
    return any(re.search(r'\b' + re.escape(k) + r'\b', text_low) for k in keywords)


def extract_passeport(text):
    """Détecte la présence d'un passeport."""
    text_low = text.lower()
    keywords = ['passeport', 'passport', 'جواز سفر']
    return any(re.search(r'\b' + re.escape(k) + r'\b', text_low) for k in keywords)


def extract_vehicule(text):
    """Détecte si le candidat est véhiculé."""
    text_low = text.lower()
    keywords = ['véhicule personnel', 'vehicule personnel', 'véhiculé', 'vehicule', 'voiture personnelle', 'own vehicle', 'own car', 'personal vehicle', 'car owner', 'سيارة شخصية']
    return any(re.search(r'\b' + re.escape(k) + r'\b', text_low) for k in keywords)


# ==========================================
# 4. SECTIONS
# ==========================================
SECTIONS_KEYWORDS = {
    'experiences': [
        'experience', 'expériences', 'expérience', 'experiences',
        'parcours professionnel', 'parcours pro', 'emploi', 'work experience',
        'professional experience', 'carrière', 'carriere'
    ],
    'formations': [
        'formation', 'formations', 'education', 'études', 'etudes',
        'diplôme', 'diplome', 'diplômes', 'diplomes', 'cursus', 'scolarité'
    ],
    'competences': [
        'compétence', 'competence', 'compétences', 'competences',
        'skills', 'aptitudes', 'savoir-faire', 'technical skills', 'informatique'
    ],
    'langues': [
        'langue', 'langues', 'languages', 'language', 'compétences linguistiques'
    ],
    # Section volontairement non exploitée : sert juste de délimiteur pour ne pas polluer
    # la section précédente (ex: compétences) avec le contenu "Centres d'intérêt".
    'interets': [
        'centres d', 'centre d', 'hobbies', 'loisirs', "centres d'intérêt", "centres d'interet",
    ],
}


def find_sections(text):
    lines = text.split("\n")
    sections = {}
    current_section = None
    current_content = []

    for line in lines:
        line_clean = line.strip().lower()
        if not line_clean:
            if current_section:
                current_content.append("")
            continue

        matched_section = None
        for section_name, keywords in SECTIONS_KEYWORDS.items():
            for keyword in keywords:
                if line_clean.startswith(keyword) and len(line_clean) < 60:
                    matched_section = section_name
                    break
            if matched_section:
                break

        if matched_section:
            if current_section:
                sections[current_section] = "\n".join(current_content).strip()
            current_section = matched_section
            current_content = []
        elif current_section:
            current_content.append(line)

    if current_section:
        sections[current_section] = "\n".join(current_content).strip()

    return sections


# ==========================================
# 5. COMPÉTENCES & LANGUES
# ==========================================
def extract_competences(text_section):
    if not text_section:
        return ""
    cleaned = re.sub(r'[•●○■□▪▫◦‣⁃]', ',', text_section)
    cleaned = re.sub(r'[-–—]\s', ',', cleaned)
    cleaned = re.sub(r'\n+', ',', cleaned)
    items = re.split(r'[,;]', cleaned)
    competences = [item.strip() for item in items if item.strip() and len(item.strip()) > 1]
    seen = set()
    result = []
    for c in competences:
        c_low = c.lower()
        if c_low not in seen and len(c) <= 80:
            seen.add(c_low)
            result.append(c)
        if len(result) >= 30:
            break
    return ", ".join(result)


def extract_langues(text_section):
    if not text_section:
        return ""
    langues_connues = [
        'français', 'francais', 'french', 'arabe', 'arabic',
        'anglais', 'english', 'espagnol', 'spanish', 'allemand', 'german',
        'italien', 'italian', 'tamazight', 'berbère', 'kabyle', 'chinois',
        'turc', 'russe', 'portugais'
    ]
    found = []
    text_low = text_section.lower()
    for langue in langues_connues:
        if langue in text_low:
            langue_clean = langue.capitalize()
            if langue_clean in [f.split(' (')[0] for f in found]:
                continue
            pattern = rf'{langue}\s*[:\-]?\s*([a-zà-ÿ0-9\s\(\)/.,]{{0,60}})'
            m = re.search(pattern, text_low)
            niveau = ""
            if m:
                niveau_raw = m.group(1).strip()
                for keyword in ['maternelle', 'natif', 'native', 'bilingue', 'courant', 'avancé', 'avance', 'fluent', 'intermédiaire', 'intermediaire', 'débutant', 'debutant', 'bon niveau', 'notions', 'scolaire', 'c1', 'c2', 'b1', 'b2', 'a1', 'a2']:
                    if keyword in niveau_raw:
                        niveau = f" ({keyword.capitalize()})"
                        break
            found.append(langue_clean + niveau)
    return ", ".join(found)


# ==========================================
# 6. EXPÉRIENCES (PATTERN AMÉLIORÉ)
# ==========================================
# Format très répandu dans les CV algériens : "Du 01/11/2021 au 04/2024 : Chef de projet chez FIELDCORE"
DATE_HEADER_PATTERN = re.compile(r'^du\s+(.+?)\s+au\s+(.+?)\s*:\s*(.*)$', re.IGNORECASE)


def _parse_titre_entreprise(texte):
    """Sépare 'Chef de projet chez FIELDCORE, GE...' en (titre, entreprise)."""
    for separateur in (r'\bchez\b', r'\bà\b', r'\bau sein de\b'):
        parts = re.split(separateur, texte, maxsplit=1, flags=re.IGNORECASE)
        if len(parts) == 2:
            return parts[0].strip(" :-"), parts[1].strip(" .")
    return texte.strip(" :-"), ""


def _extract_experiences_du_au(lines):
    """Détecte le format 'Du <date> au <date> : Titre chez Entreprise'.
    Le PDF coupe parfois cette ligne en deux (ex: "...: Organisateur de conditionnement"
    puis "chez ACG SIM." sur la ligne suivante) — on recolle les lignes de continuation
    (celles qui ne commencent pas par une puce) avant de considérer que la description commence."""
    indices = [i for i, l in enumerate(lines) if DATE_HEADER_PATTERN.match(l)]
    if not indices:
        return []
    experiences = []
    for idx, start in enumerate(indices):
        end = indices[idx + 1] if idx + 1 < len(indices) else len(lines)
        block = lines[start:end]
        m = DATE_HEADER_PATTERN.match(block[0])
        date_debut_raw, date_fin_raw, reste = m.groups()
        reste = reste.strip()

        # Recolle les lignes suivantes tant qu'elles ne sont pas des puces de description
        i = 1
        while i < len(block) and not block[i].lstrip().startswith(('-', '•', '*')):
            reste = f"{reste} {block[i].strip()}".strip()
            i += 1

        titre, entreprise = _parse_titre_entreprise(reste)

        # Format fréquent : le titre est sur la ligne D'AVANT ("Chef de projet\nDu ... au ... : chez FIELDCORE")
        # plutôt que sur la ligne de date elle-même (qui ne contient alors que "chez Entreprise").
        if not titre and start > 0:
            ligne_precedente = lines[start - 1].strip()
            if (
                ligne_precedente
                and len(ligne_precedente) < 100
                and not DATE_HEADER_PATTERN.match(ligne_precedente)
                and not ligne_precedente.lstrip().startswith(('-', '•', '*'))
            ):
                titre = ligne_precedente
                # Cette ligne avait été ajoutée à tort comme puce finale de l'expérience précédente : on la retire.
                if experiences:
                    bogus = f"- {ligne_precedente}"
                    if experiences[-1]["description"].endswith(bogus):
                        experiences[-1]["description"] = experiences[-1]["description"][: -len(bogus)].rstrip("\n")

        date_fin_present = bool(re.search(r"aujourd'?hui|présent|present|en cours", date_fin_raw, re.IGNORECASE))
        description_lines = block[i:]
        description = "\n".join(
            l if l.lstrip().startswith(('-', '•', '*')) else f"- {l}"
            for l in description_lines
        )
        if titre:
            experiences.append({
                "titre_poste": titre[:200],
                "entreprise": entreprise[:200],
                "date_debut_raw": date_debut_raw.strip(),
                "date_fin_raw": "Aujourd'hui" if date_fin_present else date_fin_raw.strip(),
                "description": description[:3000],
            })
    return experiences[:10]


def extract_experiences_list(text_section):
    """
    Extrait les expériences en détectant les titres de postes
    (lignes en majuscules ou contenant des mots-clés métier).
    Chaque expérience commence par un titre, suivi de l'entreprise, des dates, des missions.
    """
    if not text_section:
        return []

    lines_du_au = [l.strip() for l in text_section.split("\n") if l.strip()]
    experiences_du_au = _extract_experiences_du_au(lines_du_au)
    if experiences_du_au:
        return experiences_du_au

    experiences = []
    annee_pattern = re.compile(r'\b(19\d{2}|20\d{2})\b')

    # Mots-clés indiquant le début d'une expérience (titres de postes typiques)
    titre_keywords = [
        'ingénieur', 'ingenieur', 'développeur', 'developpeur', 'chef',
        'directeur', 'responsable', 'manager', 'consultant', 'analyste',
        'cadre', 'technicien', 'assistant', 'comptable', 'commercial',
        'administrateur', 'gestionnaire', 'designer', 'architecte',
        'data scientist', 'expert', 'spécialiste', 'specialiste', 'coordinateur',
        'chargé', 'charge', 'opérateur', 'operateur', 'agent', 'stagiaire',
        'hote', 'hôte', 'support', 'employé', 'employe', 'collaborateur'
    ]

    lines = [l.strip() for l in text_section.split("\n") if l.strip()]

    # On identifie les indices des lignes qui sont des titres de postes
    titre_indices = []
    for i, line in enumerate(lines):
        line_low = line.lower()
        # Critères pour considérer une ligne comme titre de poste :
        # - Contient un mot-clé de poste
        # - Est en majuscules (au moins 60% des lettres en MAJUSCULE)
        # - Pas trop long
        # - Pas une simple ligne de date
        if len(line) < 5 or len(line) > 150:
            continue
        if annee_pattern.fullmatch(line.replace(" ", "").replace("-", "").replace("—", "")):
            continue

        is_mostly_upper = sum(1 for c in line if c.isupper()) >= max(2, len([c for c in line if c.isalpha()]) * 0.6)
        contains_keyword = any(k in line_low for k in titre_keywords)

        if is_mostly_upper and contains_keyword:
            titre_indices.append(i)

    # Si aucun titre détecté, on tente l'ancienne méthode (fallback)
    if not titre_indices:
        return _extract_experiences_fallback(lines, annee_pattern)

    # On découpe en blocs basés sur les indices de titres
    for idx, start in enumerate(titre_indices):
        end = titre_indices[idx + 1] if idx + 1 < len(titre_indices) else len(lines)
        block = lines[start:end]

        if not block:
            continue

        block_text = " ".join(block)
        annees_full = re.findall(r'\b(19\d{2}|20\d{2})\b', block_text)
        date_fin_present = bool(re.search(r"aujourd'?hui|présent|present|en cours", block_text, re.IGNORECASE))

        titre = block[0]
        entreprise = ""
        description_lines = []

        # La ligne suivante après le titre est généralement l'entreprise
        # (sauf si elle est une date)
        if len(block) >= 2:
            second_line = block[1]
            if not annee_pattern.fullmatch(second_line.replace(" ", "").replace("-", "").replace("—", "")):
                entreprise = second_line
                description_lines = block[2:]
            else:
                description_lines = block[1:]

        # Filtre les dates des descriptions
        description_clean = []
        for l in description_lines:
            # Ignore les lignes qui sont juste des dates/mois
            if re.match(r'^[\s,.\-/]*(janvier|février|fevrier|mars|avril|mai|juin|juillet|août|aout|septembre|octobre|novembre|décembre|decembre)?\s*\d{4}\s*[\s,.\-/àa]*(janvier|février|fevrier|mars|avril|mai|juin|juillet|août|aout|septembre|octobre|novembre|décembre|decembre)?\s*(\d{4}|aujourd\'?hui|présent|present)?\s*$', l, re.IGNORECASE):
                continue
            description_clean.append(l)

        description = "\n".join(
            l if l.lstrip().startswith(('-', '•', '*')) else f"- {l}"
            for l in description_clean
        )

        date_debut = annees_full[0] if annees_full else ""
        if date_fin_present:
            date_fin = "Aujourd'hui"
        elif len(annees_full) >= 2:
            date_fin = annees_full[-1]
        else:
            date_fin = annees_full[0] if annees_full else ""

        if titre and len(titre) < 200:
            experiences.append({
                "titre_poste": titre[:200],
                "entreprise": entreprise[:200],
                "date_debut_raw": date_debut,
                "date_fin_raw": date_fin,
                "description": description[:3000],
            })

    return experiences[:10]


def _extract_experiences_fallback(lines, annee_pattern):
    """Ancienne méthode si la détection par mots-clés échoue."""
    experiences = []
    blocks = []
    current_block = []

    for line in lines:
        if not line.strip():
            if current_block:
                blocks.append(current_block)
                current_block = []
        else:
            current_block.append(line)
    if current_block:
        blocks.append(current_block)

    for block in blocks[:10]:
        block_text = " ".join(block)
        if not annee_pattern.search(block_text):
            continue

        annees_full = re.findall(r'\b(19\d{2}|20\d{2})\b', block_text)
        date_fin_present = bool(re.search(r"aujourd'?hui|présent|present|en cours", block_text, re.IGNORECASE))

        contenu_lines = [l for l in block if not annee_pattern.fullmatch(l.replace(" ", "").replace("-", ""))]

        if not contenu_lines:
            continue

        titre = contenu_lines[0]
        entreprise = contenu_lines[1] if len(contenu_lines) >= 2 else ""
        description = "\n".join(
            l if l.lstrip().startswith(('-', '•', '*')) else f"- {l}"
            for l in contenu_lines[2:]
        ) if len(contenu_lines) > 2 else ""

        date_debut = annees_full[0] if annees_full else ""
        date_fin = "Aujourd'hui" if date_fin_present else (annees_full[-1] if len(annees_full) >= 2 else date_debut)

        if titre and date_debut:
            experiences.append({
                "titre_poste": titre[:200],
                "entreprise": entreprise[:200],
                "date_debut_raw": date_debut,
                "date_fin_raw": date_fin,
                "description": description[:3000],
            })

    return experiences[:10]


# ==========================================
# 7. FORMATIONS
# ==========================================
_MOIS_RE = r'(?:janvier|février|fevrier|mars|avril|mai|juin|juillet|août|aout|septembre|octobre|novembre|décembre|decembre)'
DATE_FORMATION_HEADER = re.compile(rf'^(?:{_MOIS_RE}\s+)?(\d{{4}})\s*:\s*(.*)$', re.IGNORECASE)


def _parse_diplome_etablissement(texte):
    """Sépare 'MBA ... à l'I.N.S.I.M d'Oran' ou 'Ingéniorat ... Université de Blida' en (diplome, etablissement)."""
    parts = re.split(r'(\bà\b|\bau sein de\b)', texte, maxsplit=1, flags=re.IGNORECASE)
    if len(parts) == 3:
        return parts[0].strip(" :-"), parts[2].strip(" .")
    m = re.search(r'\b(Université|Universite|Institut|École|Ecole|Lycée|Lycee|Centre)\b', texte, re.IGNORECASE)
    if m:
        return texte[:m.start()].strip(" :-"), texte[m.start():].strip(" .")
    return texte.strip(" :-"), ""


def _extract_formations_mois_annee(lines):
    """Détecte le format 'Février 2023 : Diplôme à Établissement' (avec ou sans mois avant l'année)."""
    indices = [i for i, l in enumerate(lines) if DATE_FORMATION_HEADER.match(l)]
    if not indices:
        return []
    formations = []
    for idx, start in enumerate(indices):
        end = indices[idx + 1] if idx + 1 < len(indices) else len(lines)
        block = lines[start:end]
        m = DATE_FORMATION_HEADER.match(block[0])
        annee, reste = m.groups()
        reste = reste.strip()

        i = 1
        while i < len(block) and not block[i].lstrip().startswith(('-', '•', '*')):
            reste = f"{reste} {block[i].strip()}".strip()
            i += 1

        diplome, etablissement = _parse_diplome_etablissement(reste)
        description_lines = block[i:]
        description = "\n".join(
            l if l.lstrip().startswith(('-', '•', '*')) else f"- {l}"
            for l in description_lines
        )
        if diplome:
            formations.append({
                "diplome": diplome[:200],
                "etablissement": etablissement[:200],
                "date_debut_raw": annee,
                "date_fin_raw": annee,
                "description": description[:3000],
            })
    return formations[:10]


def extract_formations_list(text_section):
    """
    Extrait les formations en détectant les titres de diplômes.
    """
    if not text_section:
        return []

    lines_mois_annee = [l.strip() for l in text_section.split("\n") if l.strip()]
    formations_mois_annee = _extract_formations_mois_annee(lines_mois_annee)
    if formations_mois_annee:
        return formations_mois_annee

    formations = []
    annee_pattern = re.compile(r'\b(19\d{2}|20\d{2})\b')

    # Mots-clés indiquant le début d'une formation
    diplome_keywords = [
        'diplôme', 'diplome', 'master', 'licence', 'ingénieur', 'ingenieur',
        'doctorat', 'phd', 'magistère', 'magistere', 'baccalauréat', 'baccalaureat',
        'bac', 'ts', 'technicien', 'dut', 'bts', 'certification', 'certificat',
        'formation', 'cursus', 'études', 'etudes'
    ]

    lines = [l.strip() for l in text_section.split("\n") if l.strip()]

    titre_indices = []
    for i, line in enumerate(lines):
        line_low = line.lower()
        if len(line) < 5 or len(line) > 250:
            continue
        if annee_pattern.fullmatch(line.replace(" ", "").replace("-", "").replace("—", "")):
            continue

        is_mostly_upper = sum(1 for c in line if c.isupper()) >= max(2, len([c for c in line if c.isalpha()]) * 0.6)
        contains_keyword = any(re.search(r'\b' + re.escape(k) + r'\b', line_low) for k in diplome_keywords)

        if is_mostly_upper and contains_keyword:
            titre_indices.append(i)

    # Fallback si rien trouvé
    if not titre_indices:
        return _extract_formations_fallback(lines, annee_pattern)

    for idx, start in enumerate(titre_indices):
        end = titre_indices[idx + 1] if idx + 1 < len(titre_indices) else len(lines)
        block = lines[start:end]

        if not block:
            continue

        block_text = " ".join(block)
        annees_full = re.findall(r'\b(19\d{2}|20\d{2})\b', block_text)

        diplome = block[0]
        etablissement = ""
        description_lines = []

        if len(block) >= 2:
            second_line = block[1]
            if not annee_pattern.fullmatch(second_line.replace(" ", "").replace("-", "").replace("—", "")):
                etablissement = second_line
                description_lines = block[2:]
            else:
                description_lines = block[1:]

        description_clean = [
            l for l in description_lines
            if not re.match(r'^[\s,.\-/]*\d{4}\s*[\s,.\-/à]*\d{0,4}\s*$', l)
        ]
        description = "\n".join(
            l if l.lstrip().startswith(('-', '•', '*')) else f"- {l}"
            for l in description_clean
        )

        date_debut = annees_full[0] if annees_full else ""
        date_fin = annees_full[-1] if len(annees_full) >= 2 else date_debut

        if diplome and len(diplome) < 250:
            formations.append({
                "diplome": diplome[:250],
                "etablissement": etablissement[:200],
                "date_debut_raw": date_debut,
                "date_fin_raw": date_fin,
                "description": description[:3000],
            })

    return formations[:10]


def _extract_formations_fallback(lines, annee_pattern):
    """Ancienne méthode pour formations si pas de mots-clés détectés."""
    formations = []
    blocks = []
    current_block = []

    for line in lines:
        if not line.strip():
            if current_block:
                blocks.append(current_block)
                current_block = []
        else:
            current_block.append(line)
    if current_block:
        blocks.append(current_block)

    for block in blocks[:10]:
        block_text = " ".join(block)
        if not annee_pattern.search(block_text):
            continue
        annees_full = re.findall(r'\b(19\d{2}|20\d{2})\b', block_text)
        contenu_lines = [l for l in block if not annee_pattern.fullmatch(l.replace(" ", "").replace("-", ""))]
        if not contenu_lines:
            continue

        diplome = contenu_lines[0]
        etablissement = contenu_lines[1] if len(contenu_lines) >= 2 else ""
        date_debut = annees_full[0] if annees_full else ""
        date_fin = annees_full[-1] if len(annees_full) >= 2 else date_debut

        if diplome and date_debut:
            formations.append({
                "diplome": diplome[:250],
                "etablissement": etablissement[:200],
                "date_debut_raw": date_debut,
                "date_fin_raw": date_fin,
                "description": "",
            })

    return formations[:10]

# ==========================================
# 9. PARSING VIA OLLAMA (IA LOCALE)
# ==========================================
# === CONFIGURATION GROQ (IA cloud, gratuit et rapide) ===
GROQ_MODEL = "llama-3.3-70b-versatile"
GROQ_TIMEOUT = 30  # secondes
_groq_client = None


def get_groq_client():
    """Lazy init du client Groq (chargé à la première utilisation)."""
    global _groq_client
    if _groq_client is None:
        api_key = os.getenv("GROQ_API_KEY", "")
        if not api_key:
            return None
        _groq_client = Groq(api_key=api_key)
    return _groq_client 
PROMPT_CV_COMPLET = """Tu es un expert en analyse de CV. Voici le texte brut d'un CV.

Analyse-le et extrais TOUT ce qui suit en UN SEUL objet JSON strict (rien d'autre avant/après).

1. TITRE PROFESSIONNEL : la phrase courte juste sous le nom (ex: "INGÉNIEUR LOGICIEL"). Si absent : null.

2. EXPÉRIENCES PROFESSIONNELLES (emplois, stages rémunérés — PAS les formations/diplômes/études) :
⚠️ NE METS JAMAIS DE FORMATIONS, DIPLÔMES OU ÉTUDES DEDANS ("Licence en...", "Master en...", "Bac", "Étudiant en...", cours, certifications, écoles où la personne n'a pas travaillé).
Les VRAIES expériences sont des EMPLOIS (Caissier, Ingénieur, Développeur, Stage en entreprise...).
- Pour chaque expérience : titre_poste, entreprise, date_debut_raw, date_fin_raw, description.
- Si poste actuel (pas de date de fin) : date_fin_raw = "Aujourd'hui". Garde les dates dans le format brut du CV.
- Si info manquante : "" (chaîne vide), jamais null.
- Pour "description" : NE FUSIONNE PAS les points/missions en un paragraphe. Garde chaque point sur sa propre ligne, préfixé par "- ", séparés par \n. Ne reformule pas, reste proche du texte original.
- Si aucune expérience : tableau vide [].

3. FORMATIONS ET DIPLÔMES (Master, Licence, Bac, certifications, cours suivis — PAS les emplois) :
⚠️ NE METS JAMAIS D'EMPLOIS DEDANS (postes occupés, stages en entreprise sauf stage académique avec école).
- Pour chaque formation : diplome, etablissement, date_debut_raw, date_fin_raw, description.
- Mêmes règles de "" et de description (points sur lignes séparées, préfixés "- ") que ci-dessus.

4. INFOS PERSONNELLES :
- nom_complet : nom + prénom du candidat (ex: "FILALI Zoheir"), PAS son titre professionnel. Si absent : null.
- telephone : tous les chiffres, format brut. Si absent : null.
- competences : TOUTES les compétences techniques (langages, outils, logiciels, soft skills), séparées par virgules. Si absent : null.
- langues : format "Langue:Niveau" (ex: "Arabe:Maternelle, Anglais:Avancé"). Si pas de niveau précisé : "Intermédiaire". Si absent : null.
- linkedin : URL complète du profil LinkedIn si présente. Sinon null.
- github : URL complète du profil GitHub si présente. Sinon null.
- bio : résumé percutant et professionnel du profil en 2 phrases maximum basé sur ses expériences. Sinon null.

FORMAT EXIGÉ (JSON strict, un seul objet) :
{
  "titre_professionnel": "string ou null",
  "experiences": [
    {"titre_poste": "string", "entreprise": "string", "date_debut_raw": "string", "date_fin_raw": "string", "description": "string"}
  ],
  "formations": [
    {"diplome": "string", "etablissement": "string", "date_debut_raw": "string", "date_fin_raw": "string", "description": "string"}
  ],
  "nom_complet": "string ou null",
  "telephone": "string ou null",
  "competences": "string ou null",
  "langues": "string ou null",
  "linkedin": "string ou null",
  "github": "string ou null",
  "bio": "string ou null"
}

CV À ANALYSER :
---
{cv_text}
---

RÉPONDS UNIQUEMENT AVEC L'OBJET JSON :
"""

def _call_groq(prompt, max_tokens=3000):
    """
    Appelle Groq avec un timeout.
    Retourne la string de réponse ou None en cas d'erreur.
    """
    client = get_groq_client()
    if client is None:
        logger.warning("Groq : clé API manquante, fallback regex")
        return None

    try:
        response = client.chat.completions.create(
            model=GROQ_MODEL,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.1,
            max_tokens=max_tokens,
            timeout=GROQ_TIMEOUT,
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        logger.error("Groq : erreur - %s", e)
        return None


def _extract_json_array(content):
    """Helper : nettoie et extrait un tableau JSON depuis la réponse d'Ollama."""
    if not content:
        return None

    # Nettoie les fences markdown
    content = re.sub(r"^```(?:json)?\s*", "", content)
    content = re.sub(r"\s*```$", "", content)

    # Cherche le premier [ et le dernier ]
    first = content.find("[")
    last = content.rfind("]")
    if first == -1 or last == -1:
        return None

    json_str = content[first:last + 1]
    try:
        data = json.loads(json_str)
        if isinstance(data, list):
            return data
    except json.JSONDecodeError as e:
        logger.warning("Ollama : JSON invalide - %s", e)
    return None

def _extract_json_object(content):
    """Helper : nettoie et extrait un objet JSON depuis la réponse de Groq."""
    if not content:
        return None

    content = re.sub(r"^```(?:json)?\s*", "", content)
    content = re.sub(r"\s*```$", "", content)

    first = content.find("{")
    last = content.rfind("}")
    if first == -1 or last == -1:
        return None

    json_str = content[first:last + 1]
    try:
        data = json.loads(json_str)
        if isinstance(data, dict):
            return data
    except json.JSONDecodeError as e:
        logger.warning("Groq : JSON object invalide - %s", e)
    return None

def parse_with_groq(text):
    """
    Appelle Groq UNE SEULE FOIS pour tout extraire (titre, expériences, formations, infos perso).
    Avant : 4 appels séquentiels renvoyant chacun le CV complet en entrée → 4x la consommation
    de tokens pour le même texte, ce qui déclenchait le rate-limit Groq (429) sur les CV longs.
    Retourne un dict complet, ou None si l'appel échoue entièrement.
    """
    text_truncated = text[:12000]

    result = {
        "titre_professionnel": None,
        "experiences": [],
        "formations": [],
        "nom_complet": None,
        "telephone": None,
        "competences": None,
        "langues": None,
        "linkedin": None,
        "github": None,
        "bio": None,
    }

    logger.debug("Groq : extraction complète du CV (1 appel)...")
    content = _call_groq(
        PROMPT_CV_COMPLET.replace("{cv_text}", text_truncated),
        max_tokens=8000
    )
    infos = _extract_json_object(content)
    if not infos:
        return None

    if infos.get("titre_professionnel"):
        titre = str(infos["titre_professionnel"]).strip().strip('"').strip("'")
        if titre and titre != "NON_TROUVE" and len(titre) < 150:
            result["titre_professionnel"] = titre

    experiences = infos.get("experiences")
    if isinstance(experiences, list):
        experiences_clean = [
            e for e in experiences
            if isinstance(e, dict) and e.get("titre_poste")
        ]
        result["experiences"] = experiences_clean[:15]

    formations = infos.get("formations")
    if isinstance(formations, list):
        formations_clean = [
            f for f in formations
            if isinstance(f, dict) and f.get("diplome")
        ]
        result["formations"] = formations_clean[:10]

    if infos.get("nom_complet"):
        result["nom_complet"] = str(infos["nom_complet"]).strip()
    if infos.get("telephone"):
        tel = re.sub(r'[\s.\-]', '', str(infos["telephone"]))
        tel = re.split(r'[/,;]', tel)[0]
        result["telephone"] = tel
    if infos.get("competences"):
        result["competences"] = str(infos["competences"]).strip()
    if infos.get("langues"):
        result["langues"] = str(infos["langues"]).strip()

    result["linkedin"] = infos.get("linkedin")
    result["github"] = infos.get("github")
    result["bio"] = infos.get("bio")

    return result
# ==========================================
# 8. FONCTION PRINCIPALE
# ==========================================
def parse_cv(file_path, file_name):
    text = extract_text(file_path, file_name)

    if not text or len(text.strip()) < 50:
        return {
            "success": False,
            "error": "Impossible d'extraire du texte. Le PDF est peut-être scanné (image)."
        }

    sections = find_sections(text)

    # Extraction photo
    photo = None
    name_lower = file_name.lower()
    if name_lower.endswith(".pdf"):
        photo = extract_photo_from_pdf(file_path)
    elif name_lower.endswith((".docx", ".doc")):
        photo = extract_photo_from_docx(file_path)

    # === Champs simples via regex ===
    result = {
        "success": True,
        "email": extract_email(text),
        "wilaya": extract_wilaya(text),
        "diplome": extract_diplome_max(text),
        "specialite": extract_specialite(text),
        "service_militaire": extract_service_militaire(text),
        "permis_conduire": extract_permis(text),
        "passeport_valide": extract_passeport(text),
        "vehicule_personnel": extract_vehicule(text),
        "photo": photo,
        "linkedin": None,  # Initialisation par défaut
        "github": None,    # Initialisation par défaut
        "bio": None,       # Initialisation par défaut
    }

    # === Champs complexes via Groq ===
    ai_data = parse_with_groq(text)
    methods_used = []

    # Nom complet
    if ai_data and ai_data.get("nom_complet"):
        result["nom_complet"] = ai_data["nom_complet"]
        methods_used.append("nom:ai")
    else:
        result["nom_complet"] = extract_name(text)
        methods_used.append("nom:regex")

    # Téléphone
    if ai_data and ai_data.get("telephone"):
        result["telephone"] = ai_data["telephone"]
        methods_used.append("tel:ai")
    else:
        result["telephone"] = extract_phone(text)
        methods_used.append("tel:regex")

    # Titre professionnel
    if ai_data and ai_data.get("titre_professionnel"):
        result["titre_professionnel"] = ai_data["titre_professionnel"]
        methods_used.append("titre:ai")
    else:
        result["titre_professionnel"] = extract_titre_professionnel(text)
        methods_used.append("titre:regex")

    # Compétences
    if ai_data and ai_data.get("competences"):
        result["competences"] = ai_data["competences"]
        methods_used.append("comp:ai")
    else:
        result["competences"] = extract_competences(sections.get('competences', ''))
        methods_used.append("comp:regex")

    # Langues
    if ai_data and ai_data.get("langues"):
        result["langues"] = ai_data["langues"]
        methods_used.append("lang:ai")
    else:
        result["langues"] = extract_langues(sections.get('langues', ''))
        methods_used.append("lang:regex")

    # Récupération des nouveaux attributs spécifiques à l'IA
    if ai_data:
        result["linkedin"] = ai_data.get("linkedin")
        result["github"] = ai_data.get("github")
        result["bio"] = ai_data.get("bio")

    # Expériences
    if ai_data and ai_data.get("experiences"):
        result["experiences"] = ai_data["experiences"]
        methods_used.append("exp:ai")
    else:
        result["experiences"] = extract_experiences_list(sections.get('experiences', ''))
        methods_used.append("exp:regex")

    # Formations
    if ai_data and ai_data.get("formations"):
        result["formations"] = ai_data["formations"]
        methods_used.append("form:ai")
    else:
        result["formations"] = extract_formations_list(sections.get('formations', ''))
        methods_used.append("form:regex")

    result["parsing_method"] = " | ".join(methods_used)

    return result